import datetime
import re

from docx import Document as doc
from docx.document import Document
from docx.shared import Pt
from docx.styles.style import ParagraphStyle

from config import output_file
from console_args import ArgsHandler
from word import add_hyperlink
from yt import (
    get_playlist_info,
    get_video,
    extract_timestamps,
    time_to_seconds,
    DownloadError,
    is_playlist_url
)
from log import logger

def sanitize_filename(filename: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def process_playlist() -> None:
    info = get_playlist_info(cmd_url, cookies)
    filename = args_handler.get_name() or info[0]
    filename = sanitize_filename(filename)
    heading = d.add_heading("", level=1)
    videos_data = info[1]
    logger.info(f"Processing playlist: {filename}")
    add_hyperlink(heading, filename, cmd_url)
    got_videos = 0
    for index, video_data in enumerate(videos_data):
        video_url = video_data["url"]
        video_title = video_data["title"]
        logger.info(f"{index + 1}. {video_title} ({video_url})")
        got_videos = process_playlist_video(video_url, got_videos, cookies)
    out = output_file(filename)
    d.save(out)
    logger.info(f"Output saved to: {out}")

def process_playlist_video(video_url: str, got_videos: int, cookies: str) -> int:
    try:
        video_data = get_video(video_url, cookies=cookies)
        if isinstance(video_data, dict):
            got_videos += 1
            title = video_data['title']
            url = video_data['webpage_url']
            description = video_data.get('description', '')
            paragraph = d.add_heading(f"{got_videos}. ", level=2)
            add_hyperlink(paragraph, title, url)
            for time, desc in extract_timestamps(description):
                seconds = time_to_seconds(time)
                time_url = f"{url}&t={seconds}s"
                paragraph = d.add_paragraph("")
                add_hyperlink(paragraph, time, time_url, True)
                paragraph.add_run(f" - {desc}")
        else:
            logger.warning("Skipping unavailable video")
    except DownloadError as e:
        logger.error(f"Error while loading video: {e}")
        if "Private video" in str(e):
            logger.info("Skipping private video")
    return got_videos

def process_single_video() -> None:
    try:
        video_data = get_video(cmd_url, cookies=cookies)
        if isinstance(video_data, dict):    
            title = video_data['title']
            filename = args_handler.get_name() or title
            filename = sanitize_filename(filename)
            heading = d.add_heading("", level=1)
            add_hyperlink(heading, title, cmd_url)
            url = video_data['webpage_url']
            description = video_data.get('description', '')
            for time, desc in extract_timestamps(description):
                seconds = time_to_seconds(time)
                time_url = f"{url}&t={seconds}s"
                paragraph = d.add_paragraph("")
                add_hyperlink(paragraph, time, time_url, True)
                paragraph.add_run(f" - {desc}")
            out = output_file(filename)
            d.save(out)
            logger.info(f"Output saved to: {out}")
        else:
            logger.warning("Skipping unavailable video")
    except DownloadError as e:
        logger.error(f"Error while loading video: {e}")
        if "Private video" in str(e):
            logger.info("Skipping private video")
        
if __name__ == "__main__":
    logger.info("https://github.com/Sn3ppi/yt-video-parser")
    t1 = datetime.datetime.now()
    args_handler = ArgsHandler()
    cmd_url = args_handler.get_url()
    cookies = args_handler.get_cookies()
    d = doc()
    style = d.styles["Normal"]
    if isinstance(style, ParagraphStyle):
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)
    if is_playlist_url(cmd_url):
        process_playlist()
    else:
        process_single_video()

    t2 = datetime.datetime.now()

    logger.info(f"Time elapsed: {t2 - t1}")